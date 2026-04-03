"""
Dynamic Word Template Engine
============================
主控制器：協調模板解析、LLM 分析、文件生成
用法: python -m template_engine.main --template template.docx --document input.pdf --output output.docx
"""

import argparse
import json
import sys
from pathlib import Path

from template_engine.template_parser import TemplateParser
from template_engine.llm_client import LLMClient
from template_engine.document_filler import DocumentFiller


def main():
    parser = argparse.ArgumentParser(description="Dynamic Word Template Engine")
    parser.add_argument("--template", required=True, help="Word 模板路徑 (.docx)")
    parser.add_argument("--document", required=True, help="用戶文件路徑 (PDF/txt/docx)")
    parser.add_argument("--output", required=True, help="輸出 Word 文件路徑")
    parser.add_argument("--api-key", help="API Key（或設置環境變量）")
    parser.add_argument("--provider", default="openrouter", help="LLM Provider: openrouter|deepseek|anthropic|openai")
    parser.add_argument("--debug", action="store_true", help="顯示詳細調試資訊")
    args = parser.parse_args()

    print("=" * 60)
    print("  Dynamic Word Template Engine")
    print("=" * 60)

    # ── Step 1: 解析模板結構 ──────────────────────────────────────
    print("\n[Step 1] 解析 Word 模板結構...")
    parser_engine = TemplateParser(args.template)
    template_structure = parser_engine.parse()

    if args.debug:
        print("\n模板結構：")
        print(json.dumps(template_structure, ensure_ascii=False, indent=2))

    print(f"  ✓ 發現 {len(template_structure['placeholders'])} 個 Placeholder")
    print(f"  ✓ 發現 {len(template_structure['tables'])} 個表格")
    print(f"  ✓ 發現 {len(template_structure['sections'])} 個段落區塊")

    # ── Step 2: 提取用戶文件內容 ──────────────────────────────────
    print("\n[Step 2] 提取用戶文件內容...")
    document_content = extract_document_content(args.document)
    print(f"  ✓ 提取完成，共 {len(document_content)} 字元")

    # ── Step 3: LLM 分析並提取結構化資料 ─────────────────────────
    print("\n[Step 3] 呼叫 LLM 分析文件並匹配模板欄位...")
    llm = LLMClient(api_key=args.api_key, provider=args.provider)
    extracted_data = llm.extract_data(template_structure, document_content)

    if args.debug:
        print("\nLLM 提取結果：")
        print(json.dumps(extracted_data, ensure_ascii=False, indent=2))

    filled_count = sum(1 for v in extracted_data.get("fields", {}).values() if v)
    print(f"  ✓ 成功填充 {filled_count} 個欄位")

    # ── Step 4: 填入模板並生成文件 ────────────────────────────────
    print("\n[Step 4] 填入模板生成 Word 文件...")
    filler = DocumentFiller(args.template)
    filler.fill(extracted_data, args.output)
    print(f"  ✓ 文件已生成：{args.output}")

    print("\n" + "=" * 60)
    print("  完成！")
    print("=" * 60)


def extract_document_content(document_path: str) -> str:
    """根據文件類型提取文字內容"""
    path = Path(document_path)
    ext = path.suffix.lower()

    if ext == ".txt":
        return path.read_text(encoding="utf-8")

    elif ext == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(document_path) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
        except ImportError:
            # Fallback: 用 pypdf2
            try:
                import PyPDF2
                with open(document_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    return "\n".join(page.extract_text() for page in reader.pages)
            except ImportError:
                raise ImportError("請安裝 pdfplumber: pip install pdfplumber")

    elif ext in [".docx", ".doc"]:
        from docx import Document
        doc = Document(document_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 同時提取表格內容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)
        return "\n".join(paragraphs)

    else:
        raise ValueError(f"不支援的文件格式：{ext}")


if __name__ == "__main__":
    main()
