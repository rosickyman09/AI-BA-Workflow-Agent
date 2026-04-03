"""
document_filler.py
==================
負責將 LLM 提取的 JSON 資料填入 Word 模板：
- 支援 {{placeholder}} 替換
- 支援表格 cell 填寫
- 支援重複行（動態新增行）
- 保留原有格式（字體、顏色、粗體等）
"""

import copy
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree


class DocumentFiller:
    PLACEHOLDER_RE = re.compile(r"\{\{(.+?)\}\}|\{([A-Za-z_]\w*)\}|<<(.+?)>>")

    def __init__(self, template_path: str):
        self.template_path = template_path

    def fill(self, extracted_data: dict, output_path: str):
        """主入口：填寫模板並儲存"""
        doc = Document(self.template_path)
        fields = extracted_data.get("fields", {})
        table_data = extracted_data.get("tables", {})

        # 1. 替換段落中的 placeholder
        self._fill_paragraphs(doc, fields)

        # 2. 填寫表格
        self._fill_tables(doc, fields, table_data)

        # 3. 替換頁首頁尾中的 placeholder
        self._fill_headers_footers(doc, fields)

        doc.save(output_path)

    # ──────────────────────────────────────────────────────────────
    # 段落替換
    # ──────────────────────────────────────────────────────────────
    def _fill_paragraphs(self, doc: Document, fields: dict):
        """替換所有段落中的 placeholder，保留格式"""
        for para in doc.paragraphs:
            self._replace_in_paragraph(para, fields)

    def _replace_in_paragraph(self, para, fields: dict):
        """
        替換段落中的 placeholder
        注意：Word 會把一個 placeholder 分散到多個 run，
        所以需要先合併文字，替換後再寫回
        """
        # 收集全段文字
        full_text = "".join(run.text for run in para.runs)

        # 檢查是否有 placeholder
        if not self.PLACEHOLDER_RE.search(full_text):
            return

        # 替換 placeholder
        new_text = self._replace_placeholders(full_text, fields)

        # 寫回：清空所有 run，只保留第一個 run 的格式
        if para.runs:
            # 將替換後的文字全放進第一個 run
            para.runs[0].text = new_text
            # 清空其餘 run
            for run in para.runs[1:]:
                run.text = ""

    def _replace_placeholders(self, text: str, fields: dict) -> str:
        """將文字中所有 placeholder 替換為對應值"""
        def replacer(match):
            key = match.group(1) or match.group(2) or match.group(3)
            if key:
                key = key.strip()
                value = fields.get(key)
                if value is None:
                    # 嘗試大小寫不敏感匹配
                    for k, v in fields.items():
                        if k.lower() == key.lower():
                            value = v
                            break
                return str(value) if value is not None else f"[{key}]"
            return match.group(0)

        return self.PLACEHOLDER_RE.sub(replacer, text)

    # ──────────────────────────────────────────────────────────────
    # 表格填寫
    # ──────────────────────────────────────────────────────────────
    def _fill_tables(self, doc: Document, fields: dict, table_data: dict):
        """填寫所有表格"""
        for t_idx, table in enumerate(doc.tables):
            t_key = str(t_idx)
            rows_data = table_data.get(t_key) or table_data.get(t_idx)

            if rows_data and isinstance(rows_data, list):
                # 有動態資料：填寫重複行
                self._fill_repeating_table(table, rows_data)
            else:
                # 靜態替換：直接替換 cell 中的 placeholder
                self._fill_static_table(table, fields)

    def _fill_static_table(self, table, fields: dict):
        """靜態填寫：替換表格中的 placeholder"""
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    self._replace_in_paragraph(para, fields)

    def _fill_repeating_table(self, table, rows_data: list):
        """
        動態重複行：
        1. 找到模板行（第一個有 placeholder 的 data row）
        2. 複製模板行，填入資料
        3. 移除多餘的空模板行
        """
        # 找模板行（通常係第二行，第一行係 header）
        template_row_idx = None
        for r_idx, row in enumerate(table.rows):
            cell_texts = [cell.text.strip() for cell in row.cells]
            full_text = " ".join(cell_texts)
            if self.PLACEHOLDER_RE.search(full_text):
                template_row_idx = r_idx
                break

        if template_row_idx is None:
            # 沒有找到 placeholder 行，嘗試填寫最後一行
            template_row_idx = len(table.rows) - 1

        template_row = table.rows[template_row_idx]
        template_row_xml = copy.deepcopy(template_row._tr)

        # 記住表格 XML 元素
        tbl = table._tbl

        # 移除原有模板行
        tbl.remove(template_row._tr)

        # 為每筆資料新增一行
        for row_item in rows_data:
            # 複製模板行 XML
            new_row_xml = copy.deepcopy(template_row_xml)

            # 替換每個 cell 的內容
            for cell_xml in new_row_xml.findall(f".//{qn('w:tc')}"):
                for para_xml in cell_xml.findall(f".//{qn('w:p')}"):
                    self._replace_in_paragraph_xml(para_xml, row_item)

            tbl.append(new_row_xml)

    def _replace_in_paragraph_xml(self, para_xml, fields: dict):
        """直接在 XML 層替換 placeholder（用於複製的 XML 元素）"""
        # 收集所有 run 的文字
        runs = para_xml.findall(f".//{qn('w:r')}")
        if not runs:
            return

        full_text = "".join(
            (r.find(qn("w:t")).text or "") for r in runs if r.find(qn("w:t")) is not None
        )

        if not self.PLACEHOLDER_RE.search(full_text):
            return

        new_text = self._replace_placeholders(full_text, fields)

        # 清空所有 run，只保留第一個
        if runs:
            t_elem = runs[0].find(qn("w:t"))
            if t_elem is not None:
                t_elem.text = new_text
                t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            for run in runs[1:]:
                t_elem = run.find(qn("w:t"))
                if t_elem is not None:
                    t_elem.text = ""

    # ──────────────────────────────────────────────────────────────
    # 頁首頁尾替換
    # ──────────────────────────────────────────────────────────────
    def _fill_headers_footers(self, doc: Document, fields: dict):
        """替換頁首頁尾中的 placeholder"""
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        self._replace_in_paragraph(para, fields)

            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        self._replace_in_paragraph(para, fields)
