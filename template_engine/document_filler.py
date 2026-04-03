"""
document_filler.py
==================
負責將 LLM 提取的 JSON 資料填入 Word 模板：
- 支援 {{placeholder}} 替換
- 支援表格 cell 填寫
- 支援重複行（動態新增行）
- 保留原有格式（字體、顏色、粗體等）
"""

import copy
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree


class DocumentFiller:
    PLACEHOLDER_RE = re.compile(r"\{\{(.+?)\}\}|\{([A-Za-z_]\w*)\}|<<(.+?)>>")

    def __init__(self, template_path: str):
        self.template_path = template_path

    def fill(self, extracted_data: dict, output_path: str):
        """主入口：填寫模板並儲存"""
        doc = Document(self.template_path)
        fields = extracted_data.get("fields", {})
        table_data = extracted_data.get("tables", {})

        # 1. 替換段落中的 placeholder
        self._fill_paragraphs(doc, fields)

        # 2. 填寫表格
        self._fill_tables(doc, fields, table_data)

        # 3. 替換頁首頁尾中的 placeholder
        self._fill_headers_footers(doc, fields)

        # 4. Clean up: remove any table rows that still contain unfilled
        #    placeholder markers ({{...}} or bare [ALLCAPS_KEY] patterns).
        #    This prevents duplicate/empty template rows from appearing.
        self._remove_leftover_placeholder_rows(doc)

        doc.save(output_path)

    # ──────────────────────────────────────────────────────────────
    # 段落替換
    # ──────────────────────────────────────────────────────────────
    def _fill_paragraphs(self, doc: Document, fields: dict):
        """替換所有段落中的 placeholder，保留格式"""
        for para in doc.paragraphs:
            self._replace_in_paragraph(para, fields)

    def _replace_in_paragraph(self, para, fields: dict):
        """
        替換段落中的 placeholder
        注意：Word 會把一個 placeholder 分散到多個 run，
        所以需要先合併文字，替換後再寫回
        """
        # 收集全段文字
        full_text = "".join(run.text for run in para.runs)

        # 檢查是否有 placeholder
        if not self.PLACEHOLDER_RE.search(full_text):
            return

        # 替換 placeholder
        new_text = self._replace_placeholders(full_text, fields)

        if not para.runs:
            return

        # If no unfilled markers remain, simple write-back
        if not self._HIGHLIGHT_RE.search(new_text):
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""
            return

        # Split text by unfilled placeholders and apply yellow highlighting
        segments = self._split_by_highlights(new_text)
        fmt_run_xml = copy.deepcopy(para.runs[0]._r)

        # Remove all <w:r> elements from the paragraph
        p_elem = para._p
        for r_elem in list(p_elem.findall(qn("w:r"))):
            p_elem.remove(r_elem)

        # Create new runs for each segment
        for text_part, is_highlighted in segments:
            new_r = copy.deepcopy(fmt_run_xml)
            t_elem = new_r.find(qn("w:t"))
            if t_elem is None:
                t_elem = OxmlElement("w:t")
                new_r.append(t_elem)
            t_elem.text = text_part
            t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            if is_highlighted:
                DocumentFiller._apply_yellow_highlight_run(new_r)
            p_elem.append(new_r)

    def _replace_placeholders(self, text: str, fields: dict) -> str:
        """將文字中所有 placeholder 替換為對應值"""
        def replacer(match):
            key = match.group(1) or match.group(2) or match.group(3)
            if key:
                key = key.strip()
                value = fields.get(key)
                if value is None:
                    # 嘗試大小寫不敏感匹配
                    for k, v in fields.items():
                        if k.lower() == key.lower():
                            value = v
                            break
                return str(value) if value is not None else f"[{key}]"
            return match.group(0)

        return self.PLACEHOLDER_RE.sub(replacer, text)

    # ──────────────────────────────────────────────────────────────
    # 表格填寫
    # ──────────────────────────────────────────────────────────────
    def _fill_tables(self, doc: Document, fields: dict, table_data: dict):
        """填寫所有表格"""
        for t_idx, table in enumerate(doc.tables):
            t_key = str(t_idx)
            rows_data = table_data.get(t_key) or table_data.get(t_idx)

            if rows_data and isinstance(rows_data, list):
                # 有動態資料：填寫重複行
                self._fill_repeating_table(table, rows_data)
            else:
                # 靜態替換：直接替換 cell 中的 placeholder
                self._fill_static_table(table, fields)

    def _fill_static_table(self, table, fields: dict):
        """靜態填寫：替換表格中的 placeholder"""
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    self._replace_in_paragraph(para, fields)

    def _fill_repeating_table(self, table, rows_data: list):
        """
        動態重複行：
        1. 找到所有含 placeholder 的模板行
        2. 用第一個模板行作為範本，複製並填入資料
        3. 刪除所有原始模板行（防止重複）
        """
        # 找到所有包含 placeholder 的行（跳過 header row 0）
        template_row_indices = []
        for r_idx, row in enumerate(table.rows):
            if r_idx == 0:
                continue
            cell_texts = [cell.text.strip() for cell in row.cells]
            full_text = " ".join(cell_texts)
            if self.PLACEHOLDER_RE.search(full_text):
                template_row_indices.append(r_idx)

        if not template_row_indices:
            # 沒有找到 placeholder 行，嘗試填寫最後一行
            template_row_indices = [len(table.rows) - 1]

        # Use the first template row as the prototype for cloning
        primary_idx = template_row_indices[0]
        primary_row = table.rows[primary_idx]
        template_row_xml = copy.deepcopy(primary_row._tr)
        tbl = table._tbl
        insert_before_tr = primary_row._tr

        if not rows_data:
            # No data — remove ALL template rows entirely
            for r_idx in template_row_indices:
                try:
                    tbl.remove(table.rows[r_idx]._tr)
                except ValueError:
                    pass
            return

        # Insert data rows BEFORE the first template row (preserving position),
        # then remove ALL original template rows.
        for row_item in rows_data:
            new_row_xml = copy.deepcopy(template_row_xml)
            for cell_xml in new_row_xml.findall(f".//{qn('w:tc')}"):
                for para_xml in cell_xml.findall(f".//{qn('w:p')}"):
                    self._replace_in_paragraph_xml(para_xml, row_item)
            insert_before_tr.addprevious(new_row_xml)

        # Collect the actual XML elements to remove (in case indices shifted)
        trs_to_remove = []
        for r_idx in template_row_indices:
            try:
                trs_to_remove.append(table.rows[r_idx]._tr)
            except (IndexError, ValueError):
                pass
        for tr in trs_to_remove:
            try:
                tbl.remove(tr)
            except ValueError:
                pass

    def _replace_in_paragraph_xml(self, para_xml, fields: dict):
        """直接在 XML 層替換 placeholder（用於複製的 XML 元素）"""
        # 收集所有 run 的文字
        runs = para_xml.findall(f".//{qn('w:r')}")
        if not runs:
            return

        full_text = "".join(
            (r.find(qn("w:t")).text or "") for r in runs if r.find(qn("w:t")) is not None
        )

        if not self.PLACEHOLDER_RE.search(full_text):
            return

        new_text = self._replace_placeholders(full_text, fields)

        # If no unfilled markers remain, simple write-back
        if not self._HIGHLIGHT_RE.search(new_text):
            if runs:
                t_elem = runs[0].find(qn("w:t"))
                if t_elem is not None:
                    t_elem.text = new_text
                    t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                for run in runs[1:]:
                    t_elem = run.find(qn("w:t"))
                    if t_elem is not None:
                        t_elem.text = ""
            return

        # Split by highlights and rebuild runs with yellow background
        segments = self._split_by_highlights(new_text)
        fmt_run = copy.deepcopy(runs[0])

        # Remove all existing runs
        for r in list(runs):
            r.getparent().remove(r)

        # Create new runs for each segment
        for text_part, is_highlighted in segments:
            new_r = copy.deepcopy(fmt_run)
            t_elem = new_r.find(qn("w:t"))
            if t_elem is None:
                t_elem = OxmlElement("w:t")
                new_r.append(t_elem)
            t_elem.text = text_part
            t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            if is_highlighted:
                DocumentFiller._apply_yellow_highlight_run(new_r)
            para_xml.append(new_r)

    # ──────────────────────────────────────────────────────────────
    # 頁首頁尾替換
    # ──────────────────────────────────────────────────────────────
    def _fill_headers_footers(self, doc: Document, fields: dict):
        """替換頁首頁尾中的 placeholder"""
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        self._replace_in_paragraph(para, fields)

            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        self._replace_in_paragraph(para, fields)

    # ──────────────────────────────────────────────────────────────
    # Post-fill cleanup: remove leftover placeholder rows
    # ──────────────────────────────────────────────────────────────

    # Matches {{KEY}}, {KEY}, <<KEY>>, or bare [ALLCAPS_KEY] that look like
    # unfilled template markers (but NOT normal prose like [optional]).
    _UNFILLED_RE = re.compile(
        r"\{\{.+?\}\}"           # {{PLACEHOLDER}}
        r"|<<.+?>>"              # <<PLACEHOLDER>>
        r"|\{[A-Z_][A-Z0-9_]*\}"  # {PLACEHOLDER}  (single-brace, all caps)
        r"|\[[A-Z_][A-Z0-9_]*\]"  # [PLACEHOLDER]  (brackets, all caps)
    )

    # Regex for detecting unfilled placeholders that should get yellow highlight
    _HIGHLIGHT_RE = re.compile(
        r"\[TO BE CONFIRMED\]"
        r"|\[[A-Z_][A-Z0-9_ ]*\]"    # [PLACEHOLDER] or [MULTI WORD CAPS]
        r"|\{\{.+?\}\}"              # {{PLACEHOLDER}}
        r"|<<.+?>>"                   # <<PLACEHOLDER>>
        r"|\{[A-Z_][A-Z0-9_]*\}"     # {PLACEHOLDER}
    )

    def _split_by_highlights(self, text: str):
        """Split text into (text, is_highlighted) segments."""
        segments = []
        last_end = 0
        for m in self._HIGHLIGHT_RE.finditer(text):
            if m.start() > last_end:
                segments.append((text[last_end:m.start()], False))
            segments.append((m.group(), True))
            last_end = m.end()
        if last_end < len(text):
            segments.append((text[last_end:], False))
        return segments if segments else [(text, False)]

    @staticmethod
    def _apply_yellow_highlight_run(run_elem):
        """Apply yellow background highlight to a <w:r> XML element."""
        rPr = run_elem.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            run_elem.insert(0, rPr)
        hi = OxmlElement("w:highlight")
        hi.set(qn("w:val"), "yellow")
        rPr.append(hi)

    def _remove_leftover_placeholder_rows(self, doc: Document):
        """Remove any table rows that still contain unfilled placeholder tokens.

        This is the safety net that prevents duplicate / empty template rows
        from appearing in the final document.  It runs AFTER all fill passes.
        Only rows where EVERY non-empty cell still contains an unfilled marker
        are removed — rows that have a mix of real data and a stray marker are
        left alone (the marker will show as [TO BE CONFIRMED] which is fine).
        """
        for table in doc.tables:
            rows_to_remove = []
            for row in table.rows:
                cells_with_text = []
                cells_with_placeholder = []
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        cells_with_text.append(text)
                        if self._UNFILLED_RE.search(text):
                            cells_with_placeholder.append(text)

                # Skip header rows (row 0) — never remove them
                if row._tr == table.rows[0]._tr:
                    continue

                # Remove the row if it has content and ALL non-empty cells
                # contain at least one unfilled placeholder marker.
                if cells_with_text and len(cells_with_placeholder) == len(cells_with_text):
                    rows_to_remove.append(row._tr)

            tbl = table._tbl
            for tr in rows_to_remove:
                tbl.remove(tr)
