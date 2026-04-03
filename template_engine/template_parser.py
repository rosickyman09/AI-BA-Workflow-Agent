"""
template_parser.py
==================
自動解析任意 Word 模板嘅結構：
- 檢測 {{placeholder}} 標記
- 分析表格行列結構
- 識別模板係「有標記」定「純表格」模式
- 將結構轉為 JSON 描述，供 LLM 參考
"""

import re
import json
from pathlib import Path
from docx import Document
from docx.table import Table


class TemplateParser:
    # 支援多種 placeholder 格式
    PLACEHOLDER_PATTERNS = [
        r"\{\{(.+?)\}\}",    # {{欄位名}}
        r"\{(.+?)\}",        # {欄位名}
        r"\[(.+?)\]",        # [欄位名]
        r"<<(.+?)>>",        # <<欄位名>>
        r"___+",             # 下劃線空白（純空白欄位）
    ]

    def __init__(self, template_path: str):
        self.template_path = template_path
        self.doc = Document(template_path)
        self._placeholder_pattern = re.compile(
            r"\{\{(.+?)\}\}|\{([A-Za-z_]\w*)\}|<<(.+?)>>"
        )

    def parse(self) -> dict:
        """解析模板，返回完整結構描述"""
        structure = {
            "template_path": str(self.template_path),
            "mode": None,          # "placeholder" 或 "table_only"
            "placeholders": [],    # 所有發現的 placeholder 名稱
            "tables": [],          # 表格結構詳情
            "sections": [],        # 各段落區塊
            "fill_instructions": ""  # 給 LLM 的填寫指引
        }

        # 解析段落
        for para in self.doc.paragraphs:
            if para.text.strip():
                found = self._extract_placeholders(para.text)
                structure["sections"].append({
                    "type": "paragraph",
                    "text": para.text.strip(),
                    "placeholders": found
                })
                structure["placeholders"].extend(found)

        # 解析表格
        for t_idx, table in enumerate(self.doc.tables):
            table_info = self._parse_table(table, t_idx)
            structure["tables"].append(table_info)
            structure["placeholders"].extend(table_info["all_placeholders"])

        # 去重
        structure["placeholders"] = list(dict.fromkeys(structure["placeholders"]))

        # 判斷模板模式
        if structure["placeholders"]:
            structure["mode"] = "placeholder"
        else:
            structure["mode"] = "table_only"

        # 生成給 LLM 的填寫指引
        structure["fill_instructions"] = self._generate_fill_instructions(structure)

        return structure

    def _parse_table(self, table: Table, idx: int) -> dict:
        """解析單個表格的結構"""
        table_info = {
            "table_index": idx,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "header_row": [],
            "data_rows": [],
            "all_placeholders": [],
            "is_repeating": False,   # 係咪需要重複行（例如清單）
            "repeating_placeholders": []
        }

        for r_idx, row in enumerate(table.rows):
            row_data = []
            row_placeholders = []

            for cell in row.cells:
                cell_text = cell.text.strip()
                found = self._extract_placeholders(cell_text)
                row_data.append({
                    "text": cell_text,
                    "placeholders": found,
                    "is_empty": len(cell_text) == 0
                })
                row_placeholders.extend(found)
                table_info["all_placeholders"].extend(found)

            if r_idx == 0:
                table_info["header_row"] = row_data
            else:
                table_info["data_rows"].append({
                    "row_index": r_idx,
                    "cells": row_data,
                    "placeholders": row_placeholders
                })

        # 判斷係咪重複行模板（例如發票明細）
        # 如果多個 data row 有相似的 placeholder 結構，就係重複行
        if len(table_info["data_rows"]) > 1:
            first_row_ph = set(table_info["data_rows"][0]["placeholders"])
            second_row_ph = set(table_info["data_rows"][1]["placeholders"]) if len(table_info["data_rows"]) > 1 else set()
            if first_row_ph and first_row_ph == second_row_ph:
                table_info["is_repeating"] = True
                table_info["repeating_placeholders"] = list(first_row_ph)

        return table_info

    def _extract_placeholders(self, text: str) -> list:
        """從文字中提取所有 placeholder 名稱"""
        found = []
        for match in self._placeholder_pattern.finditer(text):
            name = match.group(1) or match.group(2) or match.group(3)
            if name:
                found.append(name.strip())
        return found

    def _generate_fill_instructions(self, structure: dict) -> str:
        """生成給 LLM 的詳細填寫指引"""
        lines = []

        if structure["mode"] == "placeholder":
            lines.append("此模板使用 {{欄位名}} 標記，請為每個欄位提取對應資料。")
            lines.append("\n需要填寫的欄位：")
            for ph in structure["placeholders"]:
                lines.append(f"  - {ph}")
        else:
            lines.append("此模板為純表格格式，沒有預設 placeholder 標記。")
            lines.append("請根據表格的欄標題，判斷每個欄位應填入的資料。")

        if structure["tables"]:
            lines.append(f"\n模板共有 {len(structure['tables'])} 個表格：")
            for t in structure["tables"]:
                lines.append(f"\n  表格 {t['table_index'] + 1}（{t['rows']} 行 x {t['cols']} 列）：")
                if t["header_row"]:
                    headers = [c["text"] for c in t["header_row"] if c["text"]]
                    lines.append(f"    欄標題：{' | '.join(headers)}")
                if t["is_repeating"]:
                    lines.append(f"    ⚠️ 此為重複行表格，資料應以列表格式提供")
                    lines.append(f"    重複欄位：{', '.join(t['repeating_placeholders'])}")

        return "\n".join(lines)

    def to_json(self) -> str:
        """返回 JSON 字串"""
        return json.dumps(self.parse(), ensure_ascii=False, indent=2)


if __name__ == "__main__":
    # 測試用
    import sys
    if len(sys.argv) > 1:
        parser = TemplateParser(sys.argv[1])
        print(parser.to_json())
