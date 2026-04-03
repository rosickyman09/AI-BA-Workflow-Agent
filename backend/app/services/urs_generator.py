"""URS template scanning, AI generation, and Google Drive save helpers."""

from __future__ import annotations

import io
import json
import logging
import os
import re
import subprocess
import tempfile
from datetime import datetime
from typing import Any, Dict, List
from zoneinfo import ZoneInfo

import httpx
from googleapiclient.http import MediaInMemoryUpload, MediaIoBaseDownload  # type: ignore
from sqlalchemy import text
from sqlalchemy.orm import Session

from app.services.google_drive import (
    find_or_create_folder,
    get_drive_service,
)

logger = logging.getLogger(__name__)


SUPPORTED_PLACEHOLDERS = {
    "{{PROJECT_NAME}}": "Name of the selected project",
    "{{PROJECT_DATE}}": "Document generation date (HK time)",
    "{{GENERATED_BY}}": "Name of user who generated doc",
    "{{REQUIREMENTS}}": "Extracted functional requirements",
    "{{DECISIONS}}": "Meeting decisions and outcomes",
    "{{ACTION_ITEMS}}": "Action items and owners",
    "{{SUMMARY}}": "Executive summary of all documents",
    "{{STAKEHOLDERS}}": "List of stakeholders mentioned",
    "{{OBJECTIVES}}": "Project objectives",
    "{{ASSUMPTIONS}}": "Assumptions listed",
    "{{RISKS}}": "Risk items identified",
    "{{APPROVED_DOCUMENTS}}": "List of source approved documents",
}


def scan_placeholders(content: str) -> list[str]:
    """Scan template content for placeholders."""
    pattern = r"\{\{[A-Z_]+\}\}"
    found = re.findall(pattern, content or "")
    return sorted(set(found))


def _download_file_bytes(service: Any, file_id: str) -> bytes:
    request = service.files().get_media(fileId=file_id, supportsAllDrives=True)
    buf = io.BytesIO()
    downloader = MediaIoBaseDownload(buf, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    return buf.getvalue()


def _extract_text_from_binary(content: bytes, filename: str, mime_type: str) -> str:
    mime = (mime_type or "").lower()
    name = (filename or "").lower()

    if "wordprocessingml" in mime or name.endswith(".docx"):
        try:
            from docx import Document  # type: ignore

            doc = Document(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs if p.text)
        except Exception as exc:
            logger.warning("DOCX extraction failed for %s: %s", filename, exc)

    if (
        "spreadsheetml" in mime
        or "ms-excel" in mime
        or name.endswith(".xlsx")
        or name.endswith(".xls")
    ):
        try:
            import openpyxl  # type: ignore

            wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            lines: List[str] = []
            for sheet in wb.worksheets:
                lines.append(f"=== Sheet: {sheet.title} ===")
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join(str(cell) if cell is not None else "" for cell in row).strip()
                    if row_text:
                        lines.append(row_text)
            return "\n".join(lines)
        except Exception as exc:
            logger.warning("XLSX extraction failed for %s: %s", filename, exc)

    if "pdf" in mime or name.endswith(".pdf"):
        try:
            from pypdf import PdfReader  # type: ignore

            reader = PdfReader(io.BytesIO(content))
            parts: List[str] = []
            for page in reader.pages:
                text_value = page.extract_text() or ""
                if text_value.strip():
                    parts.append(text_value.strip())
            return "\n".join(parts)
        except Exception as exc:
            logger.warning("PDF extraction failed for %s: %s", filename, exc)

    try:
        return content.decode("utf-8", errors="replace")
    except Exception:
        return ""


def extract_template_text(file_bytes: bytes, filename: str, mime_type: str = "") -> str:
    """Extract text from a template upload payload."""
    return _extract_text_from_binary(file_bytes, filename, mime_type)


async def download_file_content(file_id: str) -> str:
    """Download and extract text from a Google Drive file."""
    service = get_drive_service()
    if service is None:
        raise RuntimeError("google_drive_not_configured")

    file_meta = service.files().get(fileId=file_id, fields="mimeType,name").execute()
    mime = (file_meta.get("mimeType") or "").lower()
    name = file_meta.get("name") or ""

    # Google native files
    if "google-apps.document" in mime:
        exported = service.files().export(fileId=file_id, mimeType="text/plain").execute()
        return exported.decode("utf-8", errors="replace")
    if "google-apps.spreadsheet" in mime:
        exported = service.files().export(fileId=file_id, mimeType="text/csv").execute()
        return exported.decode("utf-8", errors="replace")
    if "google-apps.presentation" in mime:
        exported = service.files().export(fileId=file_id, mimeType="text/plain").execute()
        return exported.decode("utf-8", errors="replace")

    # Binary file types
    binary = _download_file_bytes(service, file_id)
    return _extract_text_from_binary(binary, name, mime)


def _fetch_document_with_optional_transcript(db: Session, doc_id: str):
    try:
        return db.execute(
            text(
                """
                SELECT
                    d.title,
                    d.google_drive_link,
                    d.google_drive_file_id,
                    d.doc_type,
                    dt.transcript_text,
                    dv.content AS latest_version_content
                FROM documents d
                LEFT JOIN document_transcripts dt ON dt.doc_id = d.doc_id
                LEFT JOIN LATERAL (
                    SELECT content
                    FROM document_versions
                    WHERE doc_id = d.doc_id
                    ORDER BY version_number DESC, created_at DESC
                    LIMIT 1
                ) dv ON TRUE
                WHERE d.doc_id = :doc_id
                """
            ),
            {"doc_id": doc_id},
        ).fetchone()
    except Exception:
        # Optional table can be absent in some environments.
        db.rollback()
        return db.execute(
            text(
                """
                SELECT
                    d.title,
                    d.google_drive_link,
                    d.google_drive_file_id,
                    d.doc_type,
                    NULL::text AS transcript_text,
                    dv.content AS latest_version_content
                FROM documents d
                LEFT JOIN LATERAL (
                    SELECT content
                    FROM document_versions
                    WHERE doc_id = d.doc_id
                    ORDER BY version_number DESC, created_at DESC
                    LIMIT 1
                ) dv ON TRUE
                WHERE d.doc_id = :doc_id
                """
            ),
            {"doc_id": doc_id},
        ).fetchone()


async def extract_content_from_documents(doc_ids: list[str], db: Session) -> dict:
    """Extract content from approved documents."""
    extracted: Dict[str, Any] = {
        "requirements": [],
        "decisions": [],
        "action_items": [],
        "risks": [],
        "stakeholders": [],
        "objectives": [],
        "assumptions": [],
        "summary": "",
        "raw_content": [],
    }

    for doc_id in doc_ids:
        doc = _fetch_document_with_optional_transcript(db, doc_id)
        if not doc:
            continue

        content_text = ""

        # Source 1: transcript from DB
        if doc[4]:
            content_text += str(doc[4])

        # Source 1b: latest document version content from DB
        if doc[5]:
            content_text = f"{content_text}\n{str(doc[5])}".strip()

        # Source 2: file content from Drive
        drive_file_id = doc[2]
        if drive_file_id:
            try:
                file_content = await download_file_content(str(drive_file_id))
                if file_content:
                    content_text = f"{content_text}\n{file_content}".strip()
            except Exception as exc:
                logger.warning("Drive download failed for doc %s: %s", doc_id, exc)

        extracted["raw_content"].append(
            {
                "doc_id": str(doc_id),
                "title": str(doc[0] or f"Document {str(doc_id)[:8]}"),
                "content": content_text.strip(),
            }
        )

    return extracted


def _trim_documents_for_prompt(raw_content: list[dict]) -> str:
    parts: List[str] = []
    for item in raw_content:
        title = item.get("title") or "Untitled"
        content = (item.get("content") or "").strip()
        # Keep token usage controlled while preserving useful context.
        if len(content) > 12000:
            content = content[:12000]
        parts.append(f"=== {title} ===\n{content}")
    return "\n\n".join(parts)


async def generate_with_ai(
    template_content: str,
    extracted_data: dict,
    project_name: str,
    user_name: str,
    placeholders: list[str],
    llm_provider: str = "auto",
) -> str:
    """Use AI to fill template placeholders.

    llm_provider: 'deepseek' | 'openrouter' | 'auto' (try deepseek then openrouter)
    """
    hk_now = datetime.now(ZoneInfo("Asia/Hong_Kong"))
    docs_text = _trim_documents_for_prompt(extracted_data.get("raw_content", []))
    approved_titles = ", ".join(
        str(item.get("title") or "Untitled")
        for item in extracted_data.get("raw_content", [])
    )

    system_prompt = (
        "You are a Senior Technical Business Analyst with 30 years of experience "
        "delivering enterprise IT projects across finance, logistics, healthcare, and government sectors.\n\n"
        "Your task is to produce a complete, audit-ready User Requirements Specification (URS) document "
        "by extracting information ONLY from the provided source documents and mapping it precisely "
        "into the URS template structure below.\n\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "⚠️  ABSOLUTE RULES — NEVER VIOLATE\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "1. NEVER invent, hallucinate, or assume any fact, name, number, or requirement.\n"
        "2. ONLY use information explicitly stated in the source documents.\n"
        "3. If a field cannot be filled from the source → write exactly: [TO BE CONFIRMED]\n"
        "4. NEVER leave a {{PLACEHOLDER}} marker in the output — always replace it with a value or [TO BE CONFIRMED].\n"
        "5. Every functional requirement MUST cite its source document.\n"
        "6. Maintain the EXACT section order, heading names, and table structure of the template.\n"
        "7. Do NOT merge, skip, rename, or reorder any section.\n\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "📋  PLACEHOLDER FILLING GUIDE\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "Use the following rules for each placeholder category:\n\n"
        "## COVER PAGE\n"
        "- PROJECT_NAME     → Official project name from source documents\n"
        "- PROJECT_CODE     → Project code / reference number if mentioned\n"
        "- PREPARED_BY      → Name of the document author / BA\n"
        "- REVIEWED_BY      → Name of reviewer if mentioned\n"
        "- APPROVED_BY      → Name of approver if mentioned\n"
        "- PROJECT_DATE     → Document date (today if not specified)\n\n"
        "## SECTION 1 — PROJECT OVERVIEW\n"
        "- PROJECT_BACKGROUND  → 2–4 sentences: business context, problem being solved, why this project exists\n"
        "- OBJECTIVE_1/2/3     → Specific, measurable project objectives. Use verb-first format:\n"
        "                         e.g. \"Reduce manual processing time by 40%\"\n"
        "                         Extract as many as found in source; use [TO BE CONFIRMED] if fewer than 3\n"
        "- SCOPE_IN            → Specific features, functions, or systems explicitly included\n"
        "                         List one item per bullet. Extract ALL in-scope items found.\n"
        "- SCOPE_OUT           → Items explicitly excluded or deferred. If not mentioned → [TO BE CONFIRMED]\n\n"
        "## SECTION 2 — STAKEHOLDERS & ROLES\n"
        "For each stakeholder row:\n"
        "- Extract real names and departments if mentioned in source documents\n"
        "- For roles not mentioned → fill Name and Dept as [TO BE CONFIRMED]\n"
        "- For STAKEHOLDER_ROLE row → add any additional stakeholders found beyond the fixed 5 roles\n"
        "  If no additional stakeholders → leave this row as [TO BE CONFIRMED]\n\n"
        "## SECTION 3.1 — FUNCTIONAL REQUIREMENTS\n"
        "CRITICAL: This is the most important section. Extract ALL functional requirements.\n"
        "Format each row as:\n"
        "- REQ_CATEGORY   → Single word category: Authentication / Dashboard / Reporting /\n"
        "                   Notification / Integration / Data / Workflow / UI / API / Security / Other\n"
        "- REQ_DESCRIPTION → \"System shall [verb] [object] [condition]\" format\n"
        "                    Be specific and testable. Do NOT use vague language.\n"
        "- US_PRIORITY    → Must Have / Should Have / Could Have / Won't Have (MoSCoW)\n"
        "                   Infer from urgency/importance language in source if not explicit\n"
        "- REQ_SOURCE     → Document title or section where this requirement was found\n\n"
        "Extract AS MANY requirements as found. Fill all 6 template rows minimum.\n"
        "If source has more than 6 requirements → add additional rows beyond FR-006.\n\n"
        "## SECTION 3.2 — NON-FUNCTIONAL REQUIREMENTS\n"
        "- Only the Compliance row has placeholders (NFR_REQUIREMENT, NFR_CRITERIA)\n"
        "- Extract any regulatory, compliance, or legal requirements from source\n"
        "- If none found → [TO BE CONFIRMED]\n"
        "- The Performance/Availability/Security/Usability/Scalability rows are pre-filled — do NOT change them\n\n"
        "## SECTION 4 — USER STORIES\n"
        "Format: \"As a [role], I want to [action], so that [benefit].\"\n"
        "- US_ROLE_FORMAT / US_ACTION_FORMAT / US_BENEFIT_FORMAT → fill the format example line\n"
        "- For each US row:\n"
        "  - US_ROLE    → The user role (admin / business owner / approver / manager / end user)\n"
        "  - US_STORY   → Full user story in \"As a... I want to... so that...\" format\n"
        "  - US_PRIORITY → MoSCoW priority\n"
        "- Extract from source documents. Fill all 5 rows minimum.\n"
        "- Acceptance criteria: extract from source or write testable criteria based on the story\n\n"
        "## SECTION 5 — BUSINESS RULES\n"
        "- BUSINESS_RULE_1 → Workflow rule (approval flows, state transitions, routing logic)\n"
        "- BUSINESS_RULE_2 → Access control rule (who can do what, role permissions)\n"
        "- BUSINESS_RULE_3 → Data retention rule (archiving, deletion, audit trail)\n"
        "- BUSINESS_RULE_4 → Any other business rule found in source\n"
        "- If fewer than 4 rules found → mark remaining as [TO BE CONFIRMED]\n\n"
        "## SECTION 6 — SYSTEM INTEGRATION REQUIREMENTS\n"
        "For each integration row:\n"
        "- INT_SYSTEM_1/2/3/4 → Name of the external system being integrated\n"
        "- INT_DESC_1/2/3/4   → What data is exchanged and the business purpose\n"
        "- INT_TYPE_4         → REST API / SFTP / Database Link / Message Queue / Webhook / Other\n"
        "- INT_DIRECTION_4    → Inbound / Outbound / Bidirectional\n"
        "- If no integrations mentioned → [TO BE CONFIRMED]\n\n"
        "## SECTION 7 — ASSUMPTIONS & CONSTRAINTS\n"
        "- ASSUMPTION_1/2/3 → Conditions assumed true for project success\n"
        "  e.g. \"Users will have internet access\", \"Existing SSO infrastructure will remain unchanged\"\n"
        "- CONSTRAINT_1/2/3 → Limitations on the project\n"
        "  e.g. budget ceiling, fixed go-live date, technology restrictions, regulatory deadlines\n"
        "- Extract from source. If fewer found → [TO BE CONFIRMED] for remainder\n\n"
        "## SECTION 9 — REVISION HISTORY\n"
        "- REVISION_DATE    → Today's date\n"
        "- REVISION_AUTHOR  → PREPARED_BY value\n"
        "- REVISION_CHANGES → \"Initial draft generated from approved documents\"\n\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "📐  OUTPUT FORMAT RULES\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "1. Use the EXACT template structure — same section numbers, headings, table columns\n"
        "2. Render tables as markdown: | Col1 | Col2 | Col3 |\n"
        "3. Use # for Section headings, ## for sub-sections\n"
        "4. Use - for bullet points (never * or numbers for lists)\n"
        "5. Bold key terms with **text**\n"
        "6. For [TO BE CONFIRMED] items — use exactly this text, no variations\n"
        "7. End with a ## Gap Analysis section listing all fields marked [TO BE CONFIRMED]\n"
    )

    placeholders_to_fill = sorted(set(placeholders or list(SUPPORTED_PLACEHOLDERS.keys())))
    user_prompt = f"""
## INPUTS

### 1. Source Documents (Meeting Minutes / Approved Documents):
{docs_text}

### 2. URS Template Structure:
{template_content}

---

## YOUR TASK

You are filling the URS template above using ONLY the source documents provided.

### Pre-filled values (inject these directly — do not ask):
- {{{{PROJECT_NAME}}}}  → {project_name}
- {{{{PROJECT_DATE}}}}  → {hk_now.strftime('%d %b %Y')}
- {{{{PREPARED_BY}}}}   → {user_name}
- {{{{GENERATED_BY}}}}  → {user_name}
- {{{{APPROVED_DOCUMENTS}}}} → {approved_titles}

### Placeholders to fill from source documents:
{json.dumps(placeholders_to_fill, ensure_ascii=False, indent=2)}

### Step-by-step instructions:

**STEP 1 — Read all source documents thoroughly.**
Identify: project purpose, stakeholders, requirements, constraints, integrations, risks.

**STEP 2 — Fill Section 1 (Project Overview).**
Write PROJECT_BACKGROUND as a clear 2–4 sentence summary.
Extract all objectives (OBJECTIVE_1/2/3) and scope items (SCOPE_IN/SCOPE_OUT).

**STEP 3 — Fill Section 2 (Stakeholders).**
Match names/departments from source to the 5 fixed roles.
Add extra stakeholders in the STAKEHOLDER_ROLE row if found.

**STEP 4 — Fill Section 3.1 (Functional Requirements). THIS IS THE MOST CRITICAL STEP.**
Extract EVERY functional requirement. Write each as:
"System shall [action] [object] [condition/constraint]"
Assign MoSCoW priority based on language: "must/critical/essential" → Must Have,
"should/important" → Should Have, "nice to have/optional" → Could Have.
Always cite the source document.

**STEP 5 — Fill Section 4 (User Stories).**
Convert requirements into user stories. One story per key user journey.
Format: "As a [role], I want to [capability], so that [business benefit]."

**STEP 6 — Fill Sections 5, 6, 7 (Business Rules, Integrations, Assumptions).**
Extract from source. Mark gaps as [TO BE CONFIRMED].

**STEP 7 — Output the complete filled URS document.**
Follow the template structure exactly.
End with ## Gap Analysis listing all [TO BE CONFIRMED] fields.

CRITICAL REMINDER: Replace ALL {{{{PLACEHOLDER}}}} markers. Never output a raw placeholder.
""".strip()

    openrouter_api_key = (os.environ.get("OPENROUTER_API_KEY") or "").strip()
    deepseek_api_key = (os.environ.get("DEEPSEEK_API_KEY") or "").strip()

    payload = {
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "max_tokens": 8000,
        "temperature": 0.1,
    }

    _OPENROUTER_MODEL_MAP = {
        "openrouter":       "openai/gpt-4o-mini",
        "gemini-2.5-pro":   "google/gemini-2.5-pro-preview-03-25",
        "minimax-m2.5":     "minimax/minimax-m1",
        "claude-haiku-4-5": "anthropic/claude-haiku-4-5",
    }
    errors: List[str] = []
    use_deepseek = llm_provider in ("auto", "deepseek")
    use_openrouter = llm_provider in {"auto", "openrouter", "gemini-2.5-pro", "minimax-m2.5", "claude-haiku-4-5"}
    openrouter_model_id = _OPENROUTER_MODEL_MAP.get(llm_provider, "openai/gpt-4o-mini")
    async with httpx.AsyncClient(timeout=httpx.Timeout(600.0, connect=15.0, read=600.0)) as client:
        # --- Provider 1: DeepSeek direct (primary) ---
        if deepseek_api_key and use_deepseek:
            try:
                collected_chunks: List[str] = []
                async with client.stream(
                    "POST",
                    "https://api.deepseek.com/chat/completions",
                    headers={
                        "Authorization": f"Bearer {deepseek_api_key}",
                        "Content-Type": "application/json",
                    },
                    json={
                        **payload,
                        "model": "deepseek-chat",
                        "stream": True,
                    },
                ) as stream_resp:
                    stream_resp.raise_for_status()
                    async for line in stream_resp.aiter_lines():
                        if not line.startswith("data: "):
                            continue
                        data_str = line[6:].strip()
                        if data_str == "[DONE]":
                            break
                        try:
                            chunk = json.loads(data_str)
                            delta = chunk.get("choices", [{}])[0].get("delta", {})
                            if delta.get("content"):
                                collected_chunks.append(delta["content"])
                        except (json.JSONDecodeError, IndexError, KeyError):
                            continue
                full_content = "".join(collected_chunks).strip()
                if full_content:
                    return full_content
                errors.append("deepseek: streaming returned empty content")
            except Exception as exc:
                logger.warning("DeepSeek direct generation failed: %s", exc)
                errors.append(f"deepseek: {exc}")

        # --- Provider 2: OpenRouter fallback ---
        if openrouter_api_key and use_openrouter:
            try:
                response = await client.post(
                    "https://openrouter.ai/api/v1/chat/completions",
                    headers={
                        "Authorization": f"Bearer {openrouter_api_key}",
                        "Content-Type": "application/json",
                        "HTTP-Referer": "https://ai-ba-agent.local",
                        "X-Title": "AI-BA-Agent",
                    },
                    json={
                        **payload,
                        "model": openrouter_model_id,
                        "stream": False,
                    },
                )
                response.raise_for_status()
                result = response.json()
                content = result["choices"][0]["message"]["content"]
                if content and str(content).strip():
                    return str(content).strip()
            except Exception as exc:
                logger.warning("OpenRouter generation failed: %s", exc)
                errors.append(f"openrouter: {exc}")

    if not openrouter_api_key and not deepseek_api_key:
        raise RuntimeError("No LLM provider configured. Set OPENROUTER_API_KEY or DEEPSEEK_API_KEY.")
    raise RuntimeError("LLM generation failed via all providers: " + " | ".join(errors))


def _post_process_placeholders(
    content: str,
    project_name: str,
    user_name: str,
    approved_titles: str,
    hk_date: str,
) -> str:
    """Clean up any remaining placeholder markers the LLM failed to substitute."""
    # 1. Remove "{{PLACEHOLDER}} = value" patterns first (keep only value)
    content = re.sub(r'\{\{[A-Z_]+\}\}\s*=\s*', '', content)
    # 2. Direct replacements for any remaining standalone {{...}} markers
    content = content.replace("{{PROJECT_NAME}}", project_name)
    content = content.replace("{{PROJECT_DATE}}", hk_date)
    content = content.replace("{{GENERATED_BY}}", user_name)
    content = content.replace("{{APPROVED_DOCUMENTS}}", approved_titles)
    # 3. Replace any still-remaining {{...}} with [TO BE CONFIRMED]
    content = re.sub(r'\{\{[A-Z_]+\}\}', '[TO BE CONFIRMED]', content)
    return content


def compute_placeholder_summary(
    generated_content: str,
    placeholders: list[str],
) -> dict:
    """Return which placeholders were filled vs left as [TO BE CONFIRMED]."""
    filled: list[str] = []
    unfilled: list[str] = []
    for ph in placeholders:
        key = ph.strip("{} ")
        # If the raw placeholder marker still appears → it was NOT filled by LLM
        if f"{{{{{key}}}}}" in generated_content:
            unfilled.append(key)
        else:
            filled.append(key)
    # Also scan for [TO BE CONFIRMED] occurrences as unfilled indicators
    tbc_count = generated_content.count("[TO BE CONFIRMED]")
    return {"filled": filled, "unfilled": unfilled, "to_be_confirmed_count": tbc_count}


# ---------------------------------------------------------------------------
# Markdown → formatted .docx builder
# ---------------------------------------------------------------------------

def _add_formatted_runs(paragraph, text: str):
    """Add runs with inline formatting: **bold**, [To be confirmed], FR-xxx."""
    from docx.shared import RGBColor  # type: ignore
    from docx.oxml.ns import nsdecls  # type: ignore
    from docx.oxml import parse_xml  # type: ignore

    parts = re.split(
        r'(\*\*[^*]+\*\*|\[to be confirmed\]|(?:FR|NFR|GAP)-\d+:)',
        text,
        flags=re.IGNORECASE,
    )
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif re.fullmatch(r'\[to be confirmed\]', part, re.IGNORECASE):
            run = paragraph.add_run(part)
            run.bold = True
            run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
            rPr = run._element.get_or_add_rPr()
            highlight = parse_xml(f'<w:highlight {nsdecls("w")} w:val="yellow"/>')
            rPr.append(highlight)
        elif re.match(r'(FR|NFR|GAP)-\d+:', part):
            run = paragraph.add_run(part)
            run.bold = True
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        else:
            paragraph.add_run(part)


def _highlight_tbc_in_docx(docx_bytes: bytes) -> bytes:
    """Post-process a rendered DOCX: apply yellow highlight to every '[To be confirmed]' run."""
    from docx import Document as DocxDocument  # type: ignore
    from docx.oxml.ns import nsdecls  # type: ignore
    from docx.oxml import parse_xml  # type: ignore

    _TBC_RE = re.compile(r'\[to\s+be\s+confirmed\]', re.IGNORECASE)
    _W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def _apply_yellow(run) -> None:
        rPr = run._element.get_or_add_rPr()
        for h in rPr.findall(f'{{{_W_NS}}}highlight'):
            rPr.remove(h)
        rPr.append(parse_xml(f'<w:highlight {nsdecls("w")} w:val="yellow"/>'))

    def _process_para(para) -> None:
        for run in para.runs:
            if _TBC_RE.search(run.text):
                _apply_yellow(run)

    try:
        doc = DocxDocument(io.BytesIO(docx_bytes))
        for para in doc.paragraphs:
            _process_para(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _process_para(para)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as exc:
        logger.warning("_highlight_tbc_in_docx failed: %s — returning original", exc)
        return docx_bytes


def _build_docx_table(doc, table_lines: list):
    """Build a Word table from markdown table lines."""
    from docx.shared import Pt, RGBColor  # type: ignore
    from docx.oxml.ns import nsdecls  # type: ignore
    from docx.oxml import parse_xml  # type: ignore

    rows_data = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows_data.append(cells)
    if not rows_data:
        return

    num_cols = max(len(row) for row in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass  # template may not define Table Grid — use default

    for r_idx, row_data in enumerate(rows_data):
        for c_idx in range(num_cols):
            cell = table.cell(r_idx, c_idx)
            cell_text = row_data[c_idx].strip() if c_idx < len(row_data) else ''
            cell.paragraphs[0].clear()
            _add_formatted_runs(cell.paragraphs[0], cell_text)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            if r_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="1E3A5F" w:val="clear"/>'
                )
                tcPr.append(shading)


def build_formatted_docx(content: str, template_bytes: bytes | None = None) -> bytes:
    """Convert LLM markdown output to a professionally formatted .docx file.

    If *template_bytes* is provided the template .docx is opened first so that
    all custom styles, fonts, colours and theme definitions carry over into the
    output.  The existing body content is cleared and replaced with the LLM
    generated content.
    """
    from docx import Document as DocxDocument  # type: ignore
    from docx.shared import Pt, RGBColor, Cm  # type: ignore

    if template_bytes:
        # Use the template as base — preserves all custom styles, themes, fonts
        doc = DocxDocument(io.BytesIO(template_bytes))
        # Clear all existing body content (paragraphs + tables)
        body = doc.element.body
        for child in list(body):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag in ("p", "tbl"):
                body.remove(child)
    else:
        doc = DocxDocument()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Ensure base styles exist (may already be defined in template)
    try:
        style = doc.styles['Normal']
        if not template_bytes:
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.space_before = Pt(2)
    except KeyError:
        pass

    # Heading styles (only override if not using template)
    if not template_bytes:
        for level, (size, color) in {
            1: (Pt(18), RGBColor(0x1E, 0x3A, 0x5F)),
            2: (Pt(14), RGBColor(0x1E, 0x3A, 0x5F)),
            3: (Pt(12), RGBColor(0x37, 0x41, 0x51)),
        }.items():
            try:
                h = doc.styles[f'Heading {level}']
                h.font.name = 'Arial'
                h.font.size = size
                h.font.color.rgb = color
                h.font.bold = True
            except KeyError:
                pass

    # ---------- helper: add heading safely ----------
    _heading_sizes = {1: Pt(18), 2: Pt(14), 3: Pt(12)}
    _heading_color = RGBColor(0x1E, 0x3A, 0x5F)

    def _safe_add_heading(text: str, level: int):
        """Add a heading — fall back to bold paragraph if style missing."""
        try:
            doc.add_heading(text, level=level)
        except KeyError:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = _heading_sizes.get(level, Pt(12))
            run.font.color.rgb = _heading_color
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)

    lines = content.split('\n')

    # ---------- de-duplicate consecutive identical headings ----------
    deduped: list[str] = []
    prev_heading = ""
    for ln in lines:
        s = ln.strip()
        heading_match = re.match(r'^(#{1,3})\s+(.+)', s)
        if heading_match:
            heading_text = heading_match.group(2).strip().rstrip('#').strip()
            normalised = re.sub(r'[\s*_`]+', ' ', heading_text).strip().lower()
            if normalised == prev_heading:
                continue  # skip duplicate heading
            prev_heading = normalised
        else:
            if s:
                prev_heading = ""
        deduped.append(ln)
    lines = deduped

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith('### '):
            _safe_add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith('## '):
            _safe_add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith('# '):
            _safe_add_heading(stripped[2:], level=1)
            i += 1
            continue

        # Table
        if stripped.startswith('|') and stripped.endswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                row_text = lines[i].strip()
                if not re.match(r'^\|[\s\-:|]+\|$', row_text):
                    table_lines.append(row_text)
                i += 1
            if table_lines:
                _build_docx_table(doc, table_lines)
            continue

        # Bullet list
        if stripped.startswith('- ') or stripped.startswith('* '):
            try:
                p = doc.add_paragraph(style='List Bullet')
            except KeyError:
                p = doc.add_paragraph()
            _add_formatted_runs(p, stripped[2:])
            i += 1
            continue

        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if m:
            try:
                p = doc.add_paragraph(style='List Number')
            except KeyError:
                p = doc.add_paragraph()
            _add_formatted_runs(p, m.group(2))
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^-{3,}$', stripped):
            p = doc.add_paragraph()
            run = p.add_run('\u2500' * 70)
            run.font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)
            run.font.size = Pt(6)
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        _add_formatted_runs(p, stripped)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_docx_with_template_engine(
    template_bytes: bytes,
    generated_content: str,
    project_name: str,
    user_name: str,
    llm_provider: str = "deepseek",
) -> bytes | None:
    """Fill a .docx template while preserving all formatting, tables, and layout.

    Primary path: docxtpl (Jinja2-based template rendering).
    Fallback:     template_engine.DocumentFiller (XML-level replacement).
    Returns filled .docx bytes, or None if the template has no placeholders.
    """
    # --- Phase 1: Parse template structure ---
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(template_bytes)
        tmp_path = tmp.name

    try:
        structure = None
        placeholders: list[str] = []

        try:
            from template_engine.template_parser import TemplateParser
            parser = TemplateParser(tmp_path)
            structure = parser.parse()
            placeholders = structure.get("placeholders", [])
        except ImportError:
            # Fallback: scan placeholders with regex
            from docx import Document as DocxDocument  # type: ignore
            doc = DocxDocument(io.BytesIO(template_bytes))
            full_text = "\n".join(p.text for p in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += "\n" + cell.text
            placeholders = list(dict.fromkeys(re.findall(r'\{\{(\w+)\}\}', full_text)))
            structure = {
                "mode": "placeholder" if placeholders else "table_only",
                "placeholders": placeholders,
                "tables": [],
                "fill_instructions": "",
            }

        has_placeholders = bool(placeholders)
        has_tables = bool(structure and structure.get("tables"))

        if not has_placeholders and not has_tables:
            logger.info("Template has no {{placeholders}} or tables — skipping structured fill")
            return None

        # --- Phase 2: Parse field values directly from generated_content ---
        # This runs for ALL LLM providers — no second LLM call, ever.
        import re as _re

        fields: dict[str, str] = {}

        def _extract_section(content: str, *heading_patterns: str) -> str:
            for pattern in heading_patterns:
                m = _re.search(
                    r'(?:^|\n)#{0,3}\s*' + _re.escape(pattern) + r'\s*\n+(.*?)(?=\n#|\n\d+\.|\Z)',
                    content, _re.IGNORECASE | _re.DOTALL
                )
                if m:
                    return m.group(1).strip()
            return ''

        def _extract_bullets(text: str) -> list[str]:
            return [
                _re.sub(r'^[-*•]\s*', '', line).strip()
                for line in text.splitlines()
                if _re.match(r'^[-*•]\s+', line.strip())
            ]

        # Cover page
        for label, key in [
            ('Project Code',    'PROJECT_CODE'),
            ('Reviewed By',     'REVIEWED_BY'),
            ('Approved By',     'APPROVED_BY'),
        ]:
            m = _re.search(rf'{_re.escape(label)}\s*[:\|]\s*(.+)', generated_content, _re.IGNORECASE)
            fields[key] = m.group(1).strip() if m else '[TO BE CONFIRMED]'

        # Section 1
        bg = _extract_section(generated_content, '1.1 Project Background', 'Project Background')
        fields['PROJECT_BACKGROUND'] = bg or '[TO BE CONFIRMED]'

        obj_text = _extract_section(generated_content, '1.2 Project Objectives', 'Project Objectives')
        objectives = _extract_bullets(obj_text)
        for i, key in enumerate(['OBJECTIVE_1', 'OBJECTIVE_2', 'OBJECTIVE_3']):
            fields[key] = objectives[i] if i < len(objectives) else '[TO BE CONFIRMED]'

        scope_in_text = _extract_section(generated_content, 'In Scope')
        scope_in = _extract_bullets(scope_in_text)
        fields['SCOPE_IN'] = '\n'.join(scope_in) if scope_in else '[TO BE CONFIRMED]'

        scope_out_text = _extract_section(generated_content, 'Out of Scope')
        scope_out = _extract_bullets(scope_out_text)
        fields['SCOPE_OUT'] = '\n'.join(scope_out) if scope_out else '[TO BE CONFIRMED]'

        # Section 2 Stakeholders
        stakeholder_section = _extract_section(generated_content,
            '2. Stakeholders', 'Stakeholders & Roles', 'Stakeholders and Roles')
        for role, name_key, dept_key in [
            ('Project Sponsor',  'SPONSOR_NAME',  'SPONSOR_DEPT'),
            ('Project Manager',  'PM_NAME',       'PM_DEPT'),
            ('Business Analyst', 'BA_NAME',       'BA_DEPT'),
            ('IT Lead',          'IT_LEAD_NAME',  'IT_LEAD_DEPT'),
            ('End User',         'END_USER_NAME', 'END_USER_DEPT'),
        ]:
            m = _re.search(
                rf'{_re.escape(role)}\s*\|?\s*([^|\n\[]+?)\s*\|?\s*([^|\n\[]+?)\s*\|',
                stakeholder_section, _re.IGNORECASE
            )
            fields.setdefault(name_key, m.group(1).strip() if m else '[TO BE CONFIRMED]')
            fields.setdefault(dept_key, m.group(2).strip() if m else '[TO BE CONFIRMED]')

        # Section 3.1 Functional Requirements
        fr_section = _extract_section(generated_content,
            '3.1 Functional Requirements', 'Functional Requirements')
        fr_rows = _re.findall(
            r'\|\s*FR-\d+\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|',
            fr_section
        )
        for i, (category, description, priority, source) in enumerate(fr_rows[:6]):
            fields[f'REQ_CATEGORY_{i+1}']    = category.strip() or '[TO BE CONFIRMED]'
            fields[f'REQ_DESCRIPTION_{i+1}'] = description.strip() or '[TO BE CONFIRMED]'
            fields[f'REQ_PRIORITY_{i+1}']    = priority.strip() or '[TO BE CONFIRMED]'
            fields[f'REQ_SOURCE_{i+1}']      = source.strip() or '[TO BE CONFIRMED]'
        fields['REQ_CATEGORY']    = fields.get('REQ_CATEGORY_1',    '[TO BE CONFIRMED]')
        fields['REQ_DESCRIPTION'] = fields.get('REQ_DESCRIPTION_1', '[TO BE CONFIRMED]')
        fields['REQ_SOURCE']      = fields.get('REQ_SOURCE_1',      '[TO BE CONFIRMED]')

        # Section 4 User Stories
        us_section = _extract_section(generated_content, '4. User Stories', 'User Stories')
        us_rows = _re.findall(
            r'\|\s*US-\d+\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|',
            us_section
        )
        for i, (story, criteria, priority) in enumerate(us_rows[:5]):
            fields[f'US_ROLE_{i+1}']     = story.strip() or '[TO BE CONFIRMED]'
            fields[f'US_STORY_{i+1}']    = story.strip() or '[TO BE CONFIRMED]'
            fields[f'US_PRIORITY_{i+1}'] = priority.strip() or '[TO BE CONFIRMED]'
        fields['US_ROLE']     = fields.get('US_ROLE_1',     '[TO BE CONFIRMED]')
        fields['US_STORY']    = fields.get('US_STORY_1',    '[TO BE CONFIRMED]')
        fields['US_PRIORITY'] = fields.get('US_PRIORITY_1', '[TO BE CONFIRMED]')

        us_format_m = _re.search(r'As a (.+?), I want to (.+?), so that (.+?)\.', generated_content)
        fields['US_ROLE_FORMAT']    = us_format_m.group(1).strip() if us_format_m else 'user'
        fields['US_ACTION_FORMAT']  = us_format_m.group(2).strip() if us_format_m else '[action]'
        fields['US_BENEFIT_FORMAT'] = us_format_m.group(3).strip() if us_format_m else '[benefit]'

        # Section 5 Business Rules
        br_section = _extract_section(generated_content, '5. Business Rules', 'Business Rules')
        br_rows = _re.findall(
            r'\|\s*BR-\d+\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|', br_section)
        for i, key in enumerate(['BUSINESS_RULE_1', 'BUSINESS_RULE_2', 'BUSINESS_RULE_3', 'BUSINESS_RULE_4']):
            fields[key] = br_rows[i][0].strip() if i < len(br_rows) else '[TO BE CONFIRMED]'

        # Section 6 Integrations
        int_section = _extract_section(generated_content,
            '6. System Integration', 'System Integration Requirements')
        int_rows = [r for r in _re.findall(
            r'\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|',
            int_section)
            if 'system' not in r[0].lower() and '---' not in r[0]]
        for i in range(1, 5):
            row = int_rows[i - 1] if i <= len(int_rows) else None
            fields[f'INT_SYSTEM_{i}'] = row[0].strip() if row else '[TO BE CONFIRMED]'
            fields[f'INT_DESC_{i}']   = row[3].strip() if row else '[TO BE CONFIRMED]'
        fields['INT_TYPE_4']      = int_rows[3][1].strip() if len(int_rows) > 3 else '[TO BE CONFIRMED]'
        fields['INT_DIRECTION_4'] = int_rows[3][2].strip() if len(int_rows) > 3 else '[TO BE CONFIRMED]'

        # Section 7 Assumptions & Constraints
        for section_key, keys in [
            ('7.1 Assumptions', ['ASSUMPTION_1', 'ASSUMPTION_2', 'ASSUMPTION_3']),
            ('7.2 Constraints', ['CONSTRAINT_1', 'CONSTRAINT_2', 'CONSTRAINT_3']),
        ]:
            items = _extract_bullets(_extract_section(generated_content, section_key))
            for i, key in enumerate(keys):
                fields[key] = items[i] if i < len(items) else '[TO BE CONFIRMED]'

        # NFR Compliance
        nfr_section = _extract_section(generated_content,
            '3.2 Non-Functional', 'Non-Functional Requirements')
        cm = _re.search(r'[Cc]ompliance\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|', nfr_section)
        fields['NFR_REQUIREMENT'] = cm.group(1).strip() if cm else '[TO BE CONFIRMED]'
        fields['NFR_CRITERIA']    = cm.group(2).strip() if cm else '[TO BE CONFIRMED]'

        # ── Revision history ──────────────────────────────
        hk_now = datetime.now(ZoneInfo("Asia/Hong_Kong"))
        fields['REVISION_DATE']    = hk_now.strftime('%d %b %Y')
        fields['REVISION_AUTHOR']  = user_name
        fields['REVISION_CHANGES'] = 'Initial draft generated from approved documents'

        # ── User Story format line ─────────────────────────
        us_format_m = _re.search(
            r'As a (.+?), I want to (.+?), so that (.+?)[.\n]',
            generated_content
        )
        fields['US_ROLE_FORMAT']    = us_format_m.group(1).strip() if us_format_m else 'user'
        fields['US_ACTION_FORMAT']  = us_format_m.group(2).strip() if us_format_m else '[action]'
        fields['US_BENEFIT_FORMAT'] = us_format_m.group(3).strip() if us_format_m else '[benefit]'

        # ── Core metadata — ALWAYS set these last (overwrite any parsed value) ──
        fields['PROJECT_NAME']  = project_name
        fields['project_name']  = project_name
        fields['PREPARED_BY']   = user_name
        fields['GENERATED_BY']  = user_name
        fields['generated_by']  = user_name
        fields['PROJECT_DATE']  = hk_now.strftime('%d %b %Y')
        fields['date']          = hk_now.strftime('%Y-%m-%d')
        fields['SPONSOR_DEPT']  = fields.get('SPONSOR_DEPT',   '[TO BE CONFIRMED]')
        fields['STAKEHOLDER_ROLE']            = fields.get('STAKEHOLDER_ROLE',           '[TO BE CONFIRMED]')
        fields['STAKEHOLDER_NAME']            = fields.get('STAKEHOLDER_NAME',           '[TO BE CONFIRMED]')
        fields['STAKEHOLDER_DEPT']            = fields.get('STAKEHOLDER_DEPT',           '[TO BE CONFIRMED]')
        fields['STAKEHOLDER_RESPONSIBILITY']  = fields.get('STAKEHOLDER_RESPONSIBILITY', '[TO BE CONFIRMED]')

        # ── Ensure EVERY template placeholder has a value ─────────────────────
        # Get the full list of undeclared variables directly from docxtpl
        # so we never miss any placeholder regardless of template changes
        try:
            from docxtpl import DocxTemplate as _DocxTpl
            _tmp_tpl = _DocxTpl(io.BytesIO(template_bytes))
            all_template_vars = _tmp_tpl.get_undeclared_template_variables()
        except Exception:
            all_template_vars = set(placeholders)

        for var in all_template_vars:
            if var not in fields or not fields[var]:
                # Try case-insensitive match from already-parsed fields
                matched = next(
                    (v for k, v in fields.items() if k.lower() == var.lower() and v),
                    None
                )
                fields[var] = matched or '[TO BE CONFIRMED]'

        extracted_data = {"fields": fields, "tables": {}}
        logger.info(
            "Direct parse (provider=%s) produced %d field values (template needs %d)",
            llm_provider, len(fields), len(all_template_vars),
        )

        # --- Phase 3: Render with docxtpl ---
        try:
            from docxtpl import DocxTemplate  # type: ignore
            tpl = DocxTemplate(io.BytesIO(template_bytes))

            # Use fields directly as context — all keys already normalised above
            context = {k: (str(v) if v is not None else '[TO BE CONFIRMED]')
                       for k, v in fields.items()}

            logger.info("docxtpl context keys (%d): %s",
                        len(context), sorted(context.keys()))

            tpl.render(context)
            buf = io.BytesIO()
            tpl.save(buf)
            result = buf.getvalue()
            result = _highlight_tbc_in_docx(result)
            logger.info(
                "docxtpl fill produced %d bytes (%d placeholders)",
                len(result), len(context),
            )
            return result

        except Exception as exc:
            import traceback
            logger.error(
                "docxtpl render FAILED:\n%s",
                traceback.format_exc(),
            )

        # --- Fallback: legacy DocumentFiller ---
        return _legacy_document_fill(tmp_path, extracted_data)

    except Exception as exc:
        logger.warning("Template engine fill failed completely: %s", exc)
        return None
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def _post_fill_tables(docx_bytes: bytes, table_data: dict) -> bytes | None:
    """Use DocumentFiller to fill dynamic table rows in an already-rendered DOCX."""
    try:
        from template_engine.document_filler import DocumentFiller
    except ImportError:
        return None

    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_in:
            tmp_in.write(docx_bytes)
            in_path = tmp_in.name

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_out:
            out_path = tmp_out.name

        filler = DocumentFiller(in_path)
        filler.fill({"fields": {}, "tables": table_data}, out_path)

        with open(out_path, "rb") as f:
            result = f.read()

        os.unlink(in_path)
        os.unlink(out_path)
        return result
    except Exception as exc:
        logger.warning("Table post-fill failed: %s", exc)
        return None


def _legacy_document_fill(tmp_path: str, extracted_data: dict) -> bytes | None:
    """Fallback: use the template_engine DocumentFiller directly."""
    try:
        from template_engine.document_filler import DocumentFiller

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as out_tmp:
            out_path = out_tmp.name

        filler = DocumentFiller(tmp_path)
        filler.fill(extracted_data, out_path)

        with open(out_path, "rb") as f:
            result = f.read()

        os.unlink(out_path)
        logger.info("DocumentFiller fallback produced %d bytes", len(result))
        return result
    except Exception as exc:
        logger.warning("DocumentFiller fallback also failed: %s", exc)
        return None


# ---------------------------------------------------------------------------
# DOCX → PDF conversion via LibreOffice headless
# ---------------------------------------------------------------------------


def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert DOCX bytes to PDF using LibreOffice headless.

    Raises RuntimeError if LibreOffice is not installed or conversion fails.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "input.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        env = os.environ.copy()
        env["HOME"] = tmpdir  # LibreOffice needs a writable HOME for its profile

        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--nofirststartwizard",
                "--convert-to", "pdf",
                "--outdir", tmpdir,
                docx_path,
            ],
            capture_output=True,
            text=True,
            timeout=120,
            env=env,
        )

        if result.returncode != 0:
            logger.error("LibreOffice conversion failed (rc=%d): %s", result.returncode, result.stderr)
            raise RuntimeError(f"LibreOffice PDF conversion failed: {result.stderr}")

        pdf_path = os.path.join(tmpdir, "input.pdf")
        if not os.path.exists(pdf_path):
            raise RuntimeError("LibreOffice conversion produced no output file")

        with open(pdf_path, "rb") as f:
            return f.read()


# ---------------------------------------------------------------------------
# PDF form field filling via pypdf
# ---------------------------------------------------------------------------


def fill_pdf_form_fields(
    pdf_bytes: bytes,
    generated_content: str,
    project_name: str,
    user_name: str,
    llm_provider: str = "deepseek",
) -> bytes | None:
    """Fill PDF form fields (AcroForm) using pypdf.

    Returns filled PDF bytes, or None if the PDF has no fillable form fields.
    """
    try:
        from pypdf import PdfReader, PdfWriter  # type: ignore
    except ImportError:
        logger.warning("pypdf not available for PDF form filling")
        return None

    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        form_fields = reader.get_fields()
        if not form_fields:
            logger.info("PDF template has no AcroForm fields")
            return None

        field_names = list(form_fields.keys())
        logger.info("Found %d PDF form fields: %s", len(field_names), field_names[:20])

        # Build field values from LLM extraction
        structure = {
            "mode": "placeholder",
            "placeholders": field_names,
            "tables": [],
            "fill_instructions": f"Fill these PDF form fields: {', '.join(field_names)}",
        }

        values: dict = {}
        try:
            from template_engine.llm_client import LLMClient
            norm = llm_provider.strip().lower()
            if norm in ("auto", ""):
                norm = "deepseek"
            llm = LLMClient(provider=norm)
            extracted = llm.extract_data(structure, generated_content)
            values = extracted.get("fields", {})
        except Exception as exc:
            logger.warning("LLM extraction for PDF form failed: %s", exc)

        # Inject metadata
        hk_now = datetime.now(ZoneInfo("Asia/Hong_Kong"))
        for key, val in {
            "PROJECT_NAME": project_name,
            "project_name": project_name,
            "PROJECT_DATE": hk_now.strftime("%d %b %Y"),
            "GENERATED_BY": user_name,
            "generated_by": user_name,
        }.items():
            values.setdefault(key, val)

        writer = PdfWriter()
        writer.append(reader)

        for page in writer.pages:
            writer.update_page_form_field_values(page, values)

        buf = io.BytesIO()
        writer.write(buf)
        result_bytes = buf.getvalue()
        logger.info("PDF form fill produced %d bytes", len(result_bytes))
        return result_bytes

    except Exception as exc:
        logger.warning("PDF form filling failed: %s", exc)
        return None


async def save_generated_to_drive(
    content: str,
    filename: str,
    project_name: str,
    file_format: str,
    template_drive_file_id: str | None = None,
    prebuilt_docx: bytes | None = None,
) -> dict:
    """Save generated document to Google Drive under Pending documents/[Project]/URS documents."""
    service = get_drive_service()
    if service is None:
        raise RuntimeError("google_drive_not_configured")

    root_id = (os.environ.get("GOOGLE_DRIVE_FOLDER_ID") or "").strip()
    if not root_id:
        raise RuntimeError("GOOGLE_DRIVE_FOLDER_ID is not set")

    pending_id = find_or_create_folder(service, "Pending documents", root_id)
    if not pending_id:
        raise RuntimeError("Unable to resolve Pending documents folder")
    project_id = find_or_create_folder(service, project_name, pending_id)
    if not project_id:
        raise RuntimeError(f"Unable to resolve project folder '{project_name}'")
    urs_folder_id = find_or_create_folder(service, "URS documents", project_id)
    if not urs_folder_id:
        raise RuntimeError("Unable to resolve URS documents folder")

    normalized_format = (file_format or "docx").strip().lower()
    if normalized_format not in {"docx", "txt", "xlsx", "pdf"}:
        normalized_format = "docx"

    output_name = filename
    if "." not in output_name:
        output_name = f"{output_name}.{normalized_format}"

    if normalized_format == "docx":
        if prebuilt_docx:
            # Use the pre-built .docx from the generate step (no LLM re-run)
            file_content = prebuilt_docx
            logger.info("Using pre-built .docx for Drive upload (%d bytes)", len(file_content))
        else:
            # Legacy fallback: build .docx on-the-fly
            tpl_bytes = None
            if template_drive_file_id:
                try:
                    tpl_bytes = _download_file_bytes(service, template_drive_file_id)
                except Exception as exc:
                    logger.warning("Could not download template for styling: %s", exc)
            file_content = build_formatted_docx(content, template_bytes=tpl_bytes)

        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif normalized_format == "pdf":
        if prebuilt_docx:
            file_content = prebuilt_docx
            logger.info("Using pre-built PDF for Drive upload (%d bytes)", len(file_content))
        else:
            # Build DOCX first, then convert to PDF via LibreOffice
            try:
                tpl_bytes = None
                if template_drive_file_id:
                    try:
                        tpl_bytes = _download_file_bytes(service, template_drive_file_id)
                    except Exception as exc:
                        logger.warning("Could not download template for PDF gen: %s", exc)
                docx_bytes = build_formatted_docx(content, template_bytes=tpl_bytes)
                file_content = convert_docx_to_pdf(docx_bytes)
                logger.info("Built PDF via LibreOffice for Drive upload (%d bytes)", len(file_content))
            except Exception as lo_exc:
                logger.warning("LibreOffice PDF conversion failed in save: %s — falling back to reportlab", lo_exc)
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.lib.units import cm
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

                pdf_buf = io.BytesIO()
                pdf_doc = SimpleDocTemplate(
                    pdf_buf, pagesize=A4,
                    leftMargin=2.5 * cm, rightMargin=2.5 * cm,
                    topMargin=2.5 * cm, bottomMargin=2.5 * cm,
                    title=filename,
                )
                pdf_styles = getSampleStyleSheet()
                story = []
                for line in content.splitlines():
                    stripped = line.rstrip()
                    if stripped.startswith("## "):
                        story.append(Paragraph(stripped[3:], pdf_styles["Heading2"]))
                    elif stripped.startswith("# "):
                        story.append(Paragraph(stripped[2:], pdf_styles["Heading1"]))
                    elif stripped == "":
                        story.append(Spacer(1, 0.3 * cm))
                    else:
                        text_val = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', stripped)
                        text_val = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text_val)
                        story.append(Paragraph(text_val, pdf_styles["Normal"]))
                if not story:
                    story.append(Paragraph(content, pdf_styles["Normal"]))
                pdf_doc.build(story)
                file_content = pdf_buf.getvalue()
        mime_type = "application/pdf"
    elif normalized_format == "xlsx":
        import openpyxl  # type: ignore

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Generated URS"
        for i, line in enumerate(content.splitlines() or [content], start=1):
            ws.cell(row=i, column=1, value=line)
        buf = io.BytesIO()
        wb.save(buf)
        file_content = buf.getvalue()
        mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    else:
        file_content = content.encode("utf-8")
        mime_type = "text/plain"

    metadata = {"name": output_name, "parents": [urs_folder_id]}
    media = MediaInMemoryUpload(file_content, mimetype=mime_type, resumable=False)
    uploaded = (
        service.files()
        .create(body=metadata, media_body=media, fields="id, webViewLink")
        .execute()
    )

    return {
        "file_id": uploaded.get("id", ""),
        "view_link": uploaded.get("webViewLink", ""),
        "folder_id": urs_folder_id,
    }
